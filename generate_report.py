import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import moment
from datetime import date
import pickle
from docx.oxml.ns import qn
from query import search

# function to get the site name via the url
def get_site_name(url):
    if url[0:8] == "https://":
        url = url[8:]

    if url[0:7] == "http://":
        url = url[7:]

    for i in range(len(url)-2):
        if url[i:(i+4)] == "www.":
            url = url[i+4:]
            break

    for i in range(len(url)):
        if url[i] == ".":
            url = url[:i]
            break

    return url

# function to type out the data for a search item
def type_article_data(doc, data):
    for tag in data.keys():
        paragraph = doc.add_paragraph(style="Normal")
        tag_run = paragraph.add_run(tag)
        tag_run.font.italic = True
        if tag == "URL":
            hyperlink = add_hyperlink(paragraph, data[tag], ": " + data[tag])
        else:
            data_run = paragraph.add_run(": " + data[tag])

# function to collect data from a search item
def collect_data(search_item):
    try:
        long_description = search_item["pagemap"]["metatags"][0]["og:description"]
    except KeyError:
        long_description = "N/A"
    try:
        title = search_item["pagemap"]["metatags"][0]["og:title"]
    except:
        title = search_item.get("title")
    data = {
        "Snippet": search_item.get("snippet"),
        "Long Description": long_description,
        "URL": search_item.get("link"),
    }
    try:
        org_name = search_item['pagemap']['metatags'][0]['og:site_name']
    except:
        org_name = get_site_name(data["URL"])
    return title, data, org_name

# function to add a hyperlink to a paragraph
def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

# set the specs for a style
def set_style(document, style_name, font_name='Times New Roman', color = (0x0, 0x0, 0x0), size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, indent = 0.0, space_after = 0, type="paragraph"):
    style = document.styles[style_name]
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(color[0], color[1], color[2])
    if type == "paragraph":
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), font_name)
        paragraph_format = style.paragraph_format
        paragraph_format.alignment = alignment
        paragraph_format.left_indent = Inches(indent)
        paragraph_format.space_after = Pt(space_after)

# set styles for report
def set_doc_styles(doc):
    # text
    set_style(doc, 'Default Paragraph Font', type="character", italic=False)
    set_style(doc, 'Normal', indent=0.5, space_after = 6)
    # date
    set_style(doc, 'Heading 1', bold=True, alignment = WD_ALIGN_PARAGRAPH.CENTER)
    # topic
    set_style(doc, 'Heading 2', bold=True)
    # query
    set_style(doc, 'Heading 3', bold = True, indent=0.25)
    # article title
    set_style(doc, 'Heading 4', bold = True, indent=0.5)
    return doc

# search the queries for a topic and
def type_topic(doc, topics, topic, pages):
    doc.add_heading(topic, 2)
    result_num = 0
    for subtopic in topics[topic]:
        print(f"Searching for subtopic: {topic if len(subtopic) == 0 else subtopic}...")
        doc.add_heading(topic if len(subtopic) == 0 else subtopic, 3)
        for page in range(1, pages+1):
            print(f"Scanning page {page}...")
            # get search results for that page for that topic
            search_results = search(f"{topic} {subtopic}", page, days=2)
            for i, search_item in enumerate(search_results, start=1):
                result_num += 1
                title, data, org_name = collect_data(search_item)
                paragraph = doc.add_heading(f"{org_name}: ", 4)
                title_run = paragraph.add_run(title)
                title_run.font.italic = True
                type_article_data(doc, data)

if __name__ == '__main__':
    doc = Document()

    set_doc_styles(doc)

    heading = doc.add_heading(date.today().strftime("%B %d, %Y"), 1)

    topics = {
        'Yuh-Line Niou': ['', 'BDS'],
        'Bill de Blasio': [''],
        'Mondaire Jones': [''],
        'Carlina Rivera': [''],
        'Dan Goldman': [''],
        'Jo Anne Simon': ['']
    }

    pages = 2

    for topic in topics.keys():
        print(f"Searching for topic: {topic}...")
        type_topic(doc, topics, topic, pages)

    doc.save('reports/'+date.today().strftime("%b_%d_%Y")+'_Report.docx')
